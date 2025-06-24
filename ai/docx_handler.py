import docx
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import re

def extract_text_from_docx(file_path):
    """Extract text content from a .docx file."""
    doc = Document(file_path)
    full_text = []
    
    for para in doc.paragraphs:
        full_text.append(para.text)
    
    return '\n'.join(full_text)

def create_revised_document(original_file_path, revised_file_path, revisions):
    """
    Create a revised document with track changes.
    
    Args:
        original_file_path: Path to the original .docx file
        revised_file_path: Path to save the revised .docx file
        revisions: List of revision dictionaries with original, revised, position, and reason
    """
    # Load the original document
    doc = Document(original_file_path)
    
    # Sort revisions by position
    revisions.sort(key=lambda x: x['position'])
    
    # Group revisions by paragraph
    paragraph_revisions = {}
    
    # Extract text from paragraphs
    paragraphs_text = [p.text for p in doc.paragraphs]
    
    # Map sentence positions to paragraphs
    sentence_to_paragraph = {}
    current_position = 0
    
    for i, para_text in enumerate(paragraphs_text):
        if not para_text.strip():
            continue
        sentence_to_paragraph[current_position] = i
        current_position += 1
    
    # Group revisions by paragraph
    for revision in revisions:
        position = revision['position']
        if position in sentence_to_paragraph:
            para_index = sentence_to_paragraph[position]
            if para_index not in paragraph_revisions:
                paragraph_revisions[para_index] = []
            paragraph_revisions[para_index].append(revision)
    
    # Create a new document for the revised version
    new_doc = Document()
    
    # Add a title for the revised document
    new_doc.add_heading('Revised Essay with Gemini AI', 0)
    
    # Add an introduction paragraph
    intro_para = new_doc.add_paragraph()
    intro_para.add_run('This document contains your essay with revisions suggested by Gemini AI. ')
    intro_para.add_run('Revisions are highlighted in red and include explanations. ')
    intro_para.add_run('Original text is shown with strikethrough formatting.').italic = True
    
    # Add a separator
    new_doc.add_paragraph('---')
    
    # Process each paragraph
    for i, paragraph in enumerate(doc.paragraphs):
        # Create a new paragraph in the revised document
        new_para = new_doc.add_paragraph()
        
        # If there are revisions for this paragraph
        if i in paragraph_revisions:
            para_text = paragraph.text
            last_end = 0
            
            # Sort revisions by their position in the text
            sorted_revisions = sorted(paragraph_revisions[i], 
                                     key=lambda x: para_text.find(x['original']))
            
            for revision in sorted_revisions:
                original = revision['original']
                revised = revision['revised']
                reason = revision['reason']
                
                # Find the position of the original text
                start = para_text.find(original, last_end)
                if start == -1:
                    continue
                
                end = start + len(original)
                
                # Add text before the revision
                if start > last_end:
                    new_para.add_run(para_text[last_end:start])
                
                # Add the original text with strikethrough
                original_run = new_para.add_run(original)
                original_run.font.strike = True
                original_run.font.color.rgb = RGBColor(128, 128, 128)
                
                # Add the revised text in red
                revised_run = new_para.add_run(" → " + revised + " ")
                revised_run.font.color.rgb = RGBColor(255, 0, 0)
                revised_run.bold = True
                
                # Add a comment with the reason
                comment_run = new_para.add_run(f"[{reason}]")
                comment_run.font.italic = True
                comment_run.font.color.rgb = RGBColor(0, 128, 0)
                
                last_end = end
            
            # Add any remaining text
            if last_end < len(para_text):
                new_para.add_run(para_text[last_end:])
        else:
            # If no revisions, just copy the paragraph
            new_para.add_run(paragraph.text)
    
    # Add a summary section
    new_doc.add_paragraph('---')
    summary_heading = new_doc.add_heading('Revision Summary', level=1)
    
    # Count revisions by type
    grammar_count = sum(1 for r in revisions if 'grammar' in r['reason'].lower())
    style_count = sum(1 for r in revisions if 'style' in r['reason'].lower())
    structure_count = sum(1 for r in revisions if 'structure' in r['reason'].lower())
    content_count = sum(1 for r in revisions if 'content' in r['reason'].lower())
    
    summary_para = new_doc.add_paragraph()
    summary_para.add_run(f'Total Revisions: {len(revisions)}\n')
    summary_para.add_run(f'Grammar and Punctuation: {grammar_count}\n')
    summary_para.add_run(f'Style and Clarity: {style_count}\n')
    summary_para.add_run(f'Structure and Organization: {structure_count}\n')
    summary_para.add_run(f'Content and Arguments: {content_count}\n')
    
    # Save the revised document
    new_doc.save(revised_file_path)
    
    return revised_file_path
