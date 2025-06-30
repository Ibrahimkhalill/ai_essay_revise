# import os
# import re
# import json
# from datetime import datetime
# import tempfile
# from typing import List, Dict, Tuple
# import io

# from django.http import JsonResponse, HttpResponse
# from django.views.decorators.csrf import csrf_exempt
# from django.views.decorators.http import require_POST
# from django.utils.timezone import now
# from werkzeug.utils import secure_filename
# from docx import Document
# from dotenv import load_dotenv
# from rest_framework.decorators import api_view, parser_classes
# from rest_framework.parsers import MultiPartParser
# from rest_framework.response import Response
# from rest_framework import status

# try:
#     from PyPDF2 import PdfReader
#     PYPDF2_AVAILABLE = True
# except ImportError:
#     PYPDF2_AVAILABLE = False

# # Load environment variables
# load_dotenv()

# # Configure Gemini AI
# GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')

# if not GEMINI_API_KEY:
#     print("Warning: GEMINI_API_KEY not found in environment variables. Please add it to your .env file.")

# try:
#     import google.generativeai as genai
#     genai.configure(api_key=GEMINI_API_KEY)
#     model = genai.GenerativeModel('gemini-2.5-flash')
# except Exception as e:
#     print(f"Error configuring Gemini AI: {e}")
#     model = None

# # Essay type classifications
# ALLOWED_ESSAY_TYPES = [
#     "Narrative Essay",
#     "Descriptive Essay",
#     "Expository Essay",
#     "Argumentative Essay",
#     "Persuasive Essay",
#     "Analytical Essay"
# ]

# ESSAY_TYPES = {
#     'Argumentative Essay': {
#         'keywords': ['argue', 'thesis', 'evidence', 'counterargument', 'claim', 'support', 'oppose', 'debate'],
#         'criteria': [
#             'Clear thesis statement',
#             'Strong evidence and examples',
#             'Counterargument acknowledgment',
#             'Logical flow of arguments',
#             'Source credibility check'
#         ]
#     },
#     'Narrative Essay': {
#         'keywords': ['story', 'experience', 'happened', 'remember', 'narrative', 'personal', 'journey'],
#         'criteria': [
#             'Clear narrative arc',
#             'Vivid imagery and descriptions',
#             'Dialogue quality',
#             'Character development',
#             'Chronological flow'
#         ]
#     },
#     'Analytical Essay': {
#         'keywords': ['analyze', 'literary', 'author', 'character', 'theme', 'symbolism', 'literary device'],
#         'criteria': [
#             'Present tense usage',
#             'Proper title italicization',
#             'Quote integration',
#             'Literary device identification',
#             'Theme analysis depth'
#         ]
#     },
#     'Expository Essay': {
#         'keywords': ['explain', 'inform', 'describe', 'process', 'how to', 'definition'],
#         'criteria': [
#             'Clear explanations',
#             'Logical organization',
#             'Supporting details',
#             'Objective tone',
#             'Factual accuracy'
#         ]
#     },
#     'Descriptive Essay': {
#         'keywords': ['describe', 'imagery', 'sensory', 'vivid', 'details', 'scene', 'depict'],
#         'criteria': [
#             'Vivid sensory details',
#             'Descriptive language',
#             'Clear imagery',
#             'Emotional impact',
#             'Cohesive description'
#         ]
#     },
#     'Persuasive Essay': {
#         'keywords': ['persuade', 'convince', 'argument', 'position', 'appeal', 'rhetoric', 'call to action'],
#         'criteria': [
#             'Clear position statement',
#             'Persuasive techniques (ethos, pathos, logos)',
#             'Strong evidence',
#             'Call to action',
#             'Audience engagement'
#         ]
#     }
# }

# class DocumentProcessor:
#     @staticmethod
#     def read_docx(file_path):
#         try:
#             doc = Document(file_path)
#             paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
#             return {
#                 'text': '\n\n'.join(paragraphs),
#                 'paragraphs': paragraphs,
#                 'paragraph_count': len(paragraphs)
#             }
#         except Exception as e:
#             raise Exception(f"Error reading DOCX file: {str(e)}")
    
#     @staticmethod
#     def read_txt(file_path):
#         try:
#             with open(file_path, 'r', encoding='utf-8') as file:
#                 content = file.read()
#                 paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
#                 return {
#                     'text': content,
#                     'paragraphs': paragraphs,
#                     'paragraph_count': len(paragraphs)
#                 }
#         except Exception as e:
#             raise Exception(f"Error reading TXT file: {str(e)}")

#     @staticmethod
#     def read_pdf(file):
#         if not PYPDF2_AVAILABLE:
#             raise Exception("PyPDF2 is not installed. Please install it to process PDF files.")
#         try:
#             pdf = PdfReader(file)
#             all_text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
#             paragraphs = [p.strip() for p in all_text.split('\n\n') if p.strip()]
#             return {
#                 'text': all_text,
#                 'paragraphs': paragraphs,
#                 'paragraph_count': len(paragraphs)
#             }
#         except Exception as e:
#             raise Exception(f"Error reading PDF file: {str(e)}")

#     @staticmethod
#     def extract_text_from_docx(file):
#         try:
#             doc = Document(file)
#             return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
#         except Exception as e:
#             raise Exception(f"Error extracting text from DOCX: {str(e)}")

#     @staticmethod
#     def extract_lines(file):
#         try:
#             if file.name.lower().endswith('.docx'):
#                 return DocumentProcessor.extract_text_from_docx(file).splitlines()
#             elif file.name.lower().endswith('.txt'):
#                 return file.read().decode('utf-8').splitlines()
#             elif file.name.lower().endswith('.pdf'):
#                 if not PYPDF2_AVAILABLE:
#                     raise Exception("PyPDF2 is not installed.")
#                 return DocumentProcessor.read_pdf(file)['text'].splitlines()
#             else:
#                 return []
#         except Exception as e:
#             raise Exception(f"Error extracting lines: {str(e)}")

# class EssayAnalyzer:
#     def __init__(self):
#         self.model = model
    
#     def extract_json(self, response_text: str) -> dict:
#         try:
#             json_start = response_text.find('{')
#             json_end = response_text.rfind('}') + 1
#             if json_start == -1 or json_end == 0:
#                 print("Error: No valid JSON found in response")
#                 return None
            
#             json_str = response_text[json_start:json_end].strip()
#             json_str = re.sub(r'^```json\n|```$', '', json_str, flags=re.MULTILINE)
#             return json.loads(json_str)
#         except json.JSONDecodeError as e:
#             print(f"Error parsing JSON: {e}")
#             return None
#         except Exception as e:
#             print(f"Unexpected error in extract_json: {e}")
#             return None
    
#     def classify_essay_type(self, text):
#         text_lower = text.lower()
#         scores = {}
        
#         for essay_type, data in ESSAY_TYPES.items():
#             score = sum(1 for keyword in data['keywords'] if keyword in text_lower)
#             scores[essay_type] = score
        
#         primary_type = max(scores, key=scores.get) if scores else 'Expository Essay'
#         high_scores = [t for t, s in scores.items() if s >= max(scores.values()) * 0.7] if scores else []
        
#         return {
#             'primary_type': primary_type,
#             'hybrid_types': high_scores if len(high_scores) > 1 else [],
#             'confidence': scores[primary_type] / len(ESSAY_TYPES[primary_type]['keywords']) if ESSAY_TYPES.get(primary_type, {}).get('keywords') else 0.5
#         }
    
#     def analyze_grammar_and_style(self, text):
#         if not self.model:
#             return self._fallback_analysis(text)
        
#         prompt = f"""
#         Analyze the following essay for grammar and style issues on a word-by-word basis.
#         Provide specific suggestions for improvement in JSON format:
#         {{
#             "word_issues": [
#                 {{
#                     "word": "original word",
#                     "suggestion": "suggested replacement",
#                     "type": "grammar|spelling|style",
#                     "position": {{ "start": start_character_position, "end": end_character_position }},
#                     "reason": "explanation"
#                 }}
#             ],
#             "overall_score": "score out of 100",
#             "strengths": ["list of strengths"],
#             "priority_improvements": ["top 3 improvements needed"]
#         }}
        
#         Essay text:
#         {text[:2000]}...
#         """
        
#         try:
#             response = self.model.generate_content(prompt)
#             response_text = response.text
#             json_start = response_text.find('{')
#             json_end = response_text.rfind('}') + 1
#             if json_start != -1 and json_end != -1:
#                 return json.loads(response_text[json_start:json_end])
#             return self._fallback_analysis(text)
#         except Exception as e:
#             print(f"Error in analysis: {e}")
#             return self._fallback_analysis(text)
    
#     def analyze_essay_specific_criteria(self, text, essay_type):
#         if not self.model:
#             return self._fallback_type_analysis(essay_type)
        
#         criteria = ESSAY_TYPES.get(essay_type, {}).get('criteria', [])
        
#         prompt = f"""
#         Analyze this {essay_type} essay based on these criteria: {', '.join(criteria)}.
#         Also check for source authenticity issues.
        
#         Provide analysis in JSON format:
#         {{
#             "criteria_analysis": [
#                 {{"criterion": "name", "score": "1-10", "feedback": "detailed feedback"}}
#             ],
#             "source_issues": [
#                 {{"issue": "description", "suggestion": "improvement"}}
#             ],
#             "type_specific_suggestions": ["list of suggestions specific to {essay_type} essays"]
#         }}
        
#         Essay text:
#         {text[:2000]}...
#         """
        
#         try:
#             response = self.model.generate_content(prompt)
#             response_text = response.text
#             json_start = response_text.find('{')
#             json_end = response_text.rfind('}') + 1
#             if json_start != -1 and json_end != -1:
#                 return json.loads(response_text[json_start:json_end])
#             return self._fallback_type_analysis(essay_type)
#         except Exception as e:
#             print(f"Error in type-specific analysis: {e}")
#             return self._fallback_type_analysis(essay_type)
    
#     def create_detection_prompt(self, essay_text):
#         return f"""
#         You are an expert writing assistant.
#         Analyze the following essay and detect its type (choose from: {', '.join(ALLOWED_ESSAY_TYPES)}). 
#         Return your result as a JSON object like:
#         {{
#           "essay_type": "Descriptive Essay"
#         }}
#         Essay:
#         \"\"\"
#         {essay_text}
#         \"\"\"
#         """
    
#     def create_transformation_prompt(self, essay_text, target_type):
#         return f"""
#         You are an academic writing assistant.
#         Revise the following essay to transform it into a well-written "{target_type}".
#         Make minimal edits, ensuring each change is either a single word or a single sentence.
#         Use tracked changes:
#         - Use <deletion>...</deletion> for removed text.
#         - Use <addition>...</addition> for added text.
#         Respond with a valid JSON object like:
#         {{
#           "essay_type": "{target_type}",
#           "corrected_essay": "The <addition>revised</addition> essay content..."
#         }}
#         Essay:
#         \"\"\"
#         {essay_text}
#         \"\"\"
#         """
    
#     def generate_interactive_suggestions(self, text):
#         if not self.model:
#             return self._fallback_suggestions(text)
        
#         words = text.split()
#         chunk_size = 500
#         all_suggestions = []
        
#         for i in range(0, len(words), chunk_size):
#             chunk_words = words[i:i + chunk_size]
#             chunk_text = ' '.join(chunk_words)
#             chunk_start_pos = len(' '.join(words[:i]))
            
#             prompt = f"""
#             Analyze the following essay excerpt on a word-by-word basis and provide suggestions:
#             {{
#                 "suggestions": [
#                     {{
#                         "id": "unique_id_{i}",
#                         "type": "grammar|spelling|style",
#                         "word": "original word",
#                         "suggested": "suggested replacement",
#                         "reason": "explanation",
#                         "position": {{
#                             "start": start_character_position,
#                             "end": end_character_position
#                         }},
#                         "severity": "high|medium|low"
#                     }}
#                 ]
#             }}
            
#             Essay excerpt:
#             {chunk_text}
            
#             Focus on:
#             1. Spelling errors
#             2. Grammar mistakes
#             3. Word choice improvements
#             4. Clarity and conciseness
#             """
            
#             try:
#                 response = self.model.generate_content(prompt)
#                 response_text = response.text
#                 json_start = response_text.find('{')
#                 json_end = response_text.rfind('}') + 1
#                 if json_start != -1 and json_end != -1:
#                     chunk_result = json.loads(response_text[json_start:json_end])
                    
#                     for suggestion in chunk_result.get('suggestions', []):
#                         suggestion['position']['start'] += chunk_start_pos
#                         suggestion['position']['end'] += chunk_start_pos
#                         suggestion['id'] = f"suggestion_{len(all_suggestions) + 1}"
                    
#                     all_suggestions.extend(chunk_result.get('suggestions', []))
                    
#             except Exception as e:
#                 print(f"Error processing chunk {i}: {e}")
#                 continue
        
#         return {"suggestions": all_suggestions[:20]}
    
#     def _fallback_analysis(self, text):
#         word_count = len(text.split())
#         return {
#             "word_issues": [
#                 {"word": "N/A", "suggestion": "Check API configuration", "type": "grammar", "position": {"start": 0, "end": 0}, "reason": "AI analysis unavailable"}
#             ],
#             "overall_score": "75",
#             "strengths": [
#                 f"Essay contains {word_count} words",
#                 "Content provided for analysis",
#                 "Basic essay format detected"
#             ],
#             "priority_improvements": [
#                 "Configure AI API key",
#                 "Manually review grammar",
#                 "Check word choice"
#             ]
#         }
    
#     def _fallback_type_analysis(self, essay_type):
#         return {
#             "criteria_analysis": [
#                 {"criterion": f"{essay_type} structure", "score": "7", "feedback": "Manual review recommended"},
#                 {"criterion": "Content relevance", "score": "8", "feedback": "Content appears relevant"},
#                 {"criterion": "Word choice", "score": "7", "feedback": "Basic vocabulary present"}
#             ],
#             "source_issues": [
#                 {"issue": "Unable to verify sources", "suggestion": "Check source credibility"}
#             ],
#             "type_specific_suggestions": [
#                 f"Review {essay_type} guidelines",
#                 "Ensure criteria are met",
#                 "Consider word-level review"
#             ]
#         }
    
#     def _fallback_suggestions(self, text):
#         suggestions = []
#         if "i " in text.lower():
#             suggestions.append({
#                 "id": "suggestion_1",
#                 "type": "grammar",
#                 "word": "i",
#                 "suggested": "I",
#                 "reason": "Capitalize pronoun 'I'",
#                 "position": {"start": text.lower().find("i "), "end": text.lower().find("i ") + 1},
#                 "severity": "medium"
#             })
        
#         if len(text.split()) < 50:
#             suggestions.append({
#                 "id": "suggestion_2",
#                 "type": "structure",
#                 "word": "N/A",
#                 "suggested": "Expand vocabulary",
#                 "reason": "Essay is too short",
#                 "position": {"start": 0, "end": 0},
#                 "severity": "low"
#             })
        
#         return {"suggestions": suggestions}

# class DocumentComparator:
#     @staticmethod
#     def tag_word_diff(line1: str, line2: str) -> Tuple[str, str]:
#         import difflib
#         diff = difflib.ndiff(line1.split(), line2.split())
#         tagged_line1, tagged_line2 = [], []
        
#         for word in diff:
#             if word.startswith('- '):
#                 tagged_line1.append(f"<del>{word[2:]}</del>")
#             elif word.startswith('+ '):
#                 tagged_line2.append(f"<ins>{word[2:]}</ins>")
#             elif word.startswith('  '):
#                 word_clean = word[2:]
#                 tagged_line1.append(word_clean)
#                 tagged_line2.append(word_clean)
#         return ' '.join(tagged_line1), ' '.join(tagged_line2)

# # Initialize analyzer and processor
# analyzer = EssayAnalyzer()
# doc_processor = DocumentProcessor()
# comparator = DocumentComparator()

# # Combined JSON and file upload endpoint for essay analysis
# @csrf_exempt
# @api_view(['POST'])
# @parser_classes([MultiPartParser])
# def analyze_essay(request):
#     try:
#         essay_text = None
#         filename = None
#         desired_type = request.data.get('desired_essay_type')

#         # CASE 1: JSON input (essay text + desired type)
#         if request.content_type.startswith('application/json'):
#             data = json.loads(request.body)
#             essay_text = data.get('essay_text')
#             desired_type = data.get('desired_essay_type')
#             if not essay_text or not desired_type:
#                 return Response({'error': 'Both "essay_text" and "desired_essay_type" are required in JSON'}, status=400)
#             if desired_type not in ALLOWED_ESSAY_TYPES:
#                 return Response({'error': f'Invalid desired_essay_type. Allowed types: {ALLOWED_ESSAY_TYPES}'}, status=400)

#         # CASE 2: File upload
#         elif 'file' in request.FILES:
#             file = request.FILES['file']
#             if not file.name.lower().endswith('.docx'):
#                 return Response({'error': 'Only .docx files are supported'}, status=400)
#             filename = secure_filename(file.name)
#             essay_text = doc_processor.extract_text_from_docx(file)

#         else:
#             return Response({'error': 'Either JSON data or a .docx file is required'}, status=400)

#         if not essay_text.strip():
#             return Response({'error': 'No text found in the input'}, status=400)
        
#         if len(essay_text.strip()) < 50:
#             return Response({'error': 'Essay is too short for meaningful analysis'}, status=400)

#         if not analyzer.model:
#             return Response({
#                 'error': 'AI analysis unavailable',
#                 'essay_type': 'Expository Essay',
#                 'corrected_essay': essay_text
#             }, status=503)

#         # STEP 1: Detect essay type
#         detection_prompt = analyzer.create_detection_prompt(essay_text)
#         try:
#             detection_response = analyzer.model.generate_content(detection_prompt)
#             detected_json = analyzer.extract_json(detection_response.text)

#             if not detected_json or "essay_type" not in detected_json:
#                 return Response({'error': 'Could not detect essay type'}, status=500)

#             detected_type = detected_json["essay_type"]
#             if detected_type not in ALLOWED_ESSAY_TYPES:
#                 detected_type = 'Expository Essay'

#             # STEP 2: Determine target type
#             target_type = desired_type if desired_type in ALLOWED_ESSAY_TYPES else detected_type

#             # STEP 3: Transform essay to target type
#             transformation_prompt = analyzer.create_transformation_prompt(essay_text, target_type)
#             transformation_response = analyzer.model.generate_content(transformation_prompt)
#             transformed_json = analyzer.extract_json(transformation_response.text)

#             if not transformed_json:
#                 return Response({'error': 'Essay transformation failed'}, status=500)

#             # Include additional analysis
#             classification = analyzer.classify_essay_type(essay_text)
#             general_analysis = analyzer.analyze_grammar_and_style(essay_text)
#             specific_analysis = analyzer.analyze_essay_specific_criteria(essay_text, target_type)
#             interactive_suggestions = analyzer.generate_interactive_suggestions(essay_text)

#             result = {
#                 'essay_type': transformed_json['essay_type'],
#                 'corrected_essay': transformed_json['corrected_essay'],
#                 'detected_type': detected_type,
#                 'final_type_used': target_type,
#                 'classification': classification,
#                 'analysis': {
#                     'general': general_analysis,
#                     'specific': specific_analysis
#                 },
#                 'interactive_suggestions': interactive_suggestions,
#                 'filename': filename,
#                 'timestamp': now().isoformat(),
#                 'word_count': len(essay_text.split()),
#                 'character_count': len(essay_text)
#             }

#             if filename:
#                 result['filename'] = filename

#             return Response(result, status=200)

#         except Exception as e:
#             print(f"Error in AI analysis: {e}")
#             return Response({'error': f'AI analysis failed: {str(e)}'}, status=500)

#     except Exception as e:
#         print(f"Error in essay analysis: {e}")
#         return Response({'error': f'Analysis failed: {str(e)}'}, status=500)

# @csrf_exempt
# @require_POST
# def upload_file(request):
#     try:
#         if 'file' not in request.FILES:
#             return JsonResponse({'error': 'No file uploaded'}, status=400)
        
#         uploaded_file = request.FILES['file']
#         if not uploaded_file.name.lower().endswith(('.txt', '.docx', '.pdf')):
#             return JsonResponse({'error': 'Unsupported file type. Please upload TXT, DOCX, or PDF files.'}, status=400)
        
#         filename = secure_filename(uploaded_file.name)
        
#         with tempfile.TemporaryDirectory() as tmpdir:
#             temp_file_path = os.path.join(tmpdir, filename)
            
#             with open(temp_file_path, 'wb+') as destination:
#                 for chunk in uploaded_file.chunks():
#                     destination.write(chunk)

#             content = {}
#             try:
#                 if filename.lower().endswith('.txt'):
#                     content = doc_processor.read_txt(temp_file_path)
#                 elif filename.lower().endswith('.docx'):
#                     content = doc_processor.read_docx(temp_file_path)
#                 elif filename.lower().endswith('.pdf'):
#                     with open(temp_file_path, 'rb') as f:
#                         content = doc_processor.read_pdf(f)
                
#                 return JsonResponse({
#                     'text': content.get('text', ''),
#                     'filename': filename,
#                     'paragraphs': content.get('paragraphs', []),
#                     'paragraph_count': content.get('paragraph_count', 0)
#                 })
                
#             except Exception as e:
#                 return JsonResponse({'error': f'Error processing file content: {str(e)}'}, status=400)
        
#     except Exception as e:
#         print(f"Error in file upload: {e}")
#         return JsonResponse({'error': f'File upload failed: {str(e)}'}, status=500)

# @csrf_exempt
# @require_POST
# def download_revision(request):
#     try:
#         data = json.loads(request.body)
#         final_text = data.get('final_text', '')
#         title = data.get('title', 'Revised Essay')
        
#         if not final_text:
#             return JsonResponse({'error': 'No final text provided'}, status=400)
        
#         doc = Document()
#         doc.add_heading(title, 0)
#         for paragraph in final_text.split('\n\n'):
#             if paragraph.strip():
#                 doc.add_paragraph(paragraph.strip())
        
#         doc_io = io.BytesIO()
#         doc.save(doc_io)
#         doc_io.seek(0)
        
#         response = HttpResponse(
#             content=doc_io.read(),
#             content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
#         )
#         response['Content-Disposition'] = f'attachment; filename="{title.replace(" ", "_")}_final.docx"'
#         return response
        
#     except Exception as e:
#         print(f"Error creating download: {e}")
#         return JsonResponse({'error': 'Failed to create download'}, status=500)

# @csrf_exempt
# @api_view(['POST'])
# @parser_classes([MultiPartParser])
# def compare_essays(request):
#     try:
#         if 'essay1' not in request.FILES or 'essay2' not in request.FILES:
#             return Response({'error': 'Please upload both essay1 and essay2 files'}, status=400)

#         file1 = request.FILES['essay1']
#         file2 = request.FILES['essay2']

#         if not (file1.name.lower().endswith(('.txt', '.docx', '.pdf')) and 
#                 file2.name.lower().endswith(('.txt', '.docx', '.pdf'))):
#             return Response({'error': 'Both files must be TXT, DOCX, or PDF format'}, status=400)

#         lines1 = doc_processor.extract_lines(file1)
#         lines2 = doc_processor.extract_lines(file2)

#         max_len = max(len(lines1), len(lines2))
#         comparison_results = []
#         all_text1 = []
#         all_text2 = []

#         if not model:
#             return Response({'error': 'AI analysis unavailable'}, status=503)

#         for i in range(max_len):
#             line1 = lines1[i] if i < len(lines1) else ""
#             line2 = lines2[i] if i < len(lines2) else ""

#             tagged1, tagged2 = comparator.tag_word_diff(line1, line2)
#             all_text1.append(line1)
#             all_text2.append(line2)

#             prompt = f"""
#             Compare the following lines from two essays and describe what changed.

#             Essay 1:
#             {line1}

#             Essay 2:
#             {line2}

#             Give a short, clear analysis of the differences.
#             """

#             try:
#                 response = model.generate_content(prompt)
#                 comparison_results.append({
#                     'line_number': i + 1,
#                     'essay1': tagged1,
#                     'essay2': tagged2,
#                     'analysis': response.text.strip()
#                 })
#             except Exception as e:
#                 comparison_results.append({
#                     'line_number': i + 1,
#                     'essay1': tagged1,
#                     'essay2': tagged2,
#                     'analysis': f"Error analyzing differences: {str(e)}"
#                 })

#         # Generate final summary
#         full_essay1 = "\n".join(all_text1)
#         full_essay2 = "\n".join(all_text2)
#         summary_prompt = f"""
#         You are given two essays. Provide a concise summary that highlights the key differences in tone, structure, ideas, and writing style.

#         Essay 1:
#         {full_essay1}

#         Essay 2:
#         {full_essay2}

#         Return the analysis in 3-5 bullet points.
#         """

#         try:
#             summary_response = model.generate_content(summary_prompt)
#             final_summary = summary_response.text.strip()
#         except Exception as e:
#             final_summary = f"Error generating summary: {str(e)}"

#         return Response({
#             'status': 'success',
#             'line_by_line_comparison': comparison_results,
#             'final_summary': final_summary,
#             'file1_name': secure_filename(file1.name),
#             'file2_name': secure_filename(file2.name),
#             'timestamp': now().isoformat()
#         }, status=200)

#     except Exception as e:
#         print(f"Error in essay comparison: {e}")
#         return Response({'error': f'Essay comparison failed: {str(e)}'}, status=500)


import os
import re
import json
from datetime import datetime
import tempfile
from typing import List, Dict, Tuple
import io
import difflib

from rest_framework.decorators import api_view, parser_classes
from rest_framework.parsers import MultiPartParser, JSONParser
from rest_framework.response import Response
from rest_framework import status
from django.http import HttpResponse # For sending files directly
from django.core.files.storage import default_storage # If you need to save files temporarily
from django.conf import settings # If you need MEDIA_ROOT etc.

from docx import Document
from docx.shared import RGBColor # Not used in Flask code, but kept if you plan to use it
from docx.enum.text import WD_COLOR # Not used in Flask code, but kept if you plan to use it
from dotenv import load_dotenv

# Use standard Python's tempfile for temporary file handling in Django context
# import tempfile # Already imported

try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

# Load environment variables (ensure .env file is in project root or accessible)
load_dotenv()

# Configure Gemini AI
GEMINI_API_KEY = os.getenv('GEMINI_API_KEY')

if not GEMINI_API_KEY:
    print("Warning: GEMINI_API_KEY not found in environment variables. Please add it to your .env file.")

# Initialize Gemini AI model globally once
model = None
try:
    import google.generativeai as genai
    if GEMINI_API_KEY:
        genai.configure(api_key=GEMINI_API_KEY)
        # Using a stable, publicly available model. 'gemini-2.0-flash-exp' might be experimental/internal.
        # It's safer to use 'gemini-1.5-flash' or 'gemini-pro' for general access.
        # Let's try to get gemini-1.5-flash and fallback to gemini-pro if not found.
        available_models = [m.name for m in genai.list_models()]
       
        if 'models/gemini-1.5-flash' in available_models:
            model = genai.GenerativeModel('models/gemini-1.5-flash')
           
        elif 'gemini-1.5-flash' in available_models:
            model = genai.GenerativeModel('gemini-1.5-flash')
          
        elif 'models/gemini-pro' in available_models:
            model = genai.GenerativeModel('models/gemini-pro')
           
        else:
            print("Warning: Neither 'gemini-1.5-flash' nor 'gemini-pro' found. AI features will be limited.")

except Exception as e:
    print(f"Error configuring Gemini AI: {e}")
    model = None

# Essay type classifications
ALLOWED_ESSAY_TYPES = [
    "Narrative Essay",
    "Descriptive Essay",
    "Expository Essay",
    "Argumentative Essay",
    "Persuasive Essay",
    "Analytical Essay"
]

ESSAY_TYPES = {
    'Argumentative Essay': {
        'keywords': ['argue', 'thesis', 'evidence', 'counterargument', 'claim', 'support', 'oppose', 'debate'],
        'criteria': [
            'Clear thesis statement',
            'Strong evidence and examples',
            'Counterargument acknowledgment',
            'Logical flow of arguments',
            'Source credibility check'
        ]
    },
    'Narrative Essay': {
        'keywords': ['story', 'experience', 'happened', 'remember', 'narrative', 'personal', 'journey'],
        'criteria': [
            'Clear narrative arc',
            'Vivid imagery and descriptions',
            'Dialogue quality',
            'Character development',
            'Chronological flow'
        ]
    },
    'Analytical Essay': {
        'keywords': ['analyze', 'literary', 'author', 'character', 'theme', 'symbolism', 'literary device'],
        'criteria': [
            'Present tense usage',
            'Proper title italicization',
            'Quote integration',
            'Literary device identification',
            'Theme analysis depth'
        ]
    },
    'Expository Essay': {
        'keywords': ['explain', 'inform', 'describe', 'process', 'how to', 'definition'],
        'criteria': [
            'Clear explanations',
            'Logical organization',
            'Supporting details',
            'Objective tone',
            'Factual accuracy'
        ]
    },
    'Descriptive Essay': {
        'keywords': ['describe', 'imagery', 'sensory', 'vivid', 'details', 'scene', 'depict'],
        'criteria': [
            'Vivid sensory details',
            'Descriptive language',
            'Clear imagery',
            'Emotional impact',
            'Cohesive description'
        ]
    },
    'Persuasive Essay': {
        'keywords': ['persuade', 'convince', 'argument', 'position', 'appeal', 'rhetoric', 'call to action'],
        'criteria': [
            'Clear position statement',
            'Persuasive techniques (ethos, pathos, logos)',
            'Strong evidence',
            'Call to action',
            'Audience engagement'
        ]
    }
}

class DocumentProcessor:
    @staticmethod
    def read_docx(file_object): # Changed file_path to file_object for in-memory files
        try:
            doc = Document(file_object)
            paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            return {
                'text': '\n\n'.join(paragraphs),
                'paragraphs': paragraphs,
                'paragraph_count': len(paragraphs)
            }
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {str(e)}")

    @staticmethod
    def read_txt(file_object): # Changed file_path to file_object
        try:
            content = file_object.read().decode('utf-8')
            paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
            return {
                'text': content,
                'paragraphs': paragraphs,
                'paragraph_count': len(paragraphs)
            }
        except Exception as e:
            raise Exception(f"Error reading TXT file: {str(e)}")

    @staticmethod
    def read_pdf(file_object): # Changed file to file_object
        if not PYPDF2_AVAILABLE:
            raise Exception("PyPDF2 is not installed. Please install it to process PDF files.")
        try:
            # PyPDF2 PdfReader can take a file-like object directly
            pdf = PdfReader(file_object)
            all_text = "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
            paragraphs = [p.strip() for p in all_text.split('\n\n') if p.strip()]
            return {
                'text': all_text,
                'paragraphs': paragraphs,
                'paragraph_count': len(paragraphs)
            }
        except Exception as e:
            raise Exception(f"Error reading PDF file: {str(e)}")

    @staticmethod
    def extract_text_from_docx(file_object): # Changed file to file_object
        try:
            doc = Document(file_object)
            return "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")

    @staticmethod
    def extract_lines(file_object, filename): # Added filename to get extension
        try:
            file_extension = filename.lower().split('.')[-1]
            if file_extension == 'docx':
                return DocumentProcessor.extract_text_from_docx(file_object).splitlines()
            elif file_extension == 'txt':
                return file_object.read().decode('utf-8').splitlines()
            elif file_extension == 'pdf':
                if not PYPDF2_AVAILABLE:
                    raise Exception("PyPDF2 is not installed.")
                # Reset file pointer for PDF reading if it's already been read
                file_object.seek(0)
                return DocumentProcessor.read_pdf(file_object)['text'].splitlines()
            else:
                return []
        except Exception as e:
            raise Exception(f"Error extracting lines: {str(e)}")


class EssayAnalyzer:
    def __init__(self):
        # 'model' is now a global variable, so no need for self.model here
        pass

    def sanitize_text(self, text: str) -> str:
        """Sanitize text to ensure it can be safely included in JSON."""
        if not isinstance(text, str):
            return str(text)
        # Escape quotes, newlines, and other problematic characters
        # Note: json.dumps already handles this well for string values
        # If the AI returns raw text that *needs* to be part of a JSON string,
        # then this might still be useful, but for direct JSON parsing, less so.
        text = text.replace('"', '\\"').replace('\n', '\\n').replace('\r', '\\r').replace('\t', '\\t')
        return text

    def extract_json(self, response_text: str) -> dict:
        """Extract and parse JSON from AI response, handling malformed JSON."""
        try:
            json_start = response_text.find('{')
            json_end = response_text.rfind('}') + 1
            if json_start == -1 or json_end == 0:
                print("Error: No valid JSON found in response")
                return None

            json_str = response_text[json_start:json_end].strip()
            json_str = re.sub(r'^```json\n|```$', '', json_str, flags=re.MULTILINE)
            
            # Parse JSON directly. The sanitize_text function should be applied if needed *before* JSON generation by AI
            # or if JSON string values are later inserted into another string that needs escaping.
            # json.loads handles standard JSON string unescaping.
            parsed_json = json.loads(json_str)
            return parsed_json
        except json.JSONDecodeError as e:
            print(f"Error parsing JSON: {e}")
            print(f"Problematic JSON string: {json_str[:500]}...") # Print a snippet for debugging
            return None
        except Exception as e:
            print(f"Unexpected error in extract_json: {e}")
            return None

    def classify_essay_type(self, text):
        text_lower = text.lower()
        scores = {}

        for essay_type, data in ESSAY_TYPES.items():
            score = sum(1 for keyword in data['keywords'] if keyword in text_lower)
            scores[essay_type] = score

        primary_type = max(scores, key=scores.get) if scores else 'Expository Essay'
        return primary_type

    def analyze_grammar_and_style(self, text):
        global model # Access the global model instance
        if not model:
            return self._fallback_analysis(text)

        prompt = f"""
        Analyze the following essay for grammar and style issues.
        Provide an overall score out of 100 and specific suggestions for improvement in JSON format:
        {{
            "overall_score": "score out of 100",
            "suggestions": ["suggestion 1", "suggestion 2", ...]
        }}
        
        Essay text:
        {self.sanitize_text(text[:2000])}...
        """
        
        try:
            response = model.generate_content(prompt)
            response_text = response.text
            return self.extract_json(response_text) # Use the general extract_json method
        except Exception as e:
            print(f"Error in analysis: {e}")
            return self._fallback_analysis(text)

    def create_correction_prompt(self, essay_text, essay_type):
        # Ensure essay_text is properly truncated for prompt to avoid exceeding token limits
        # The prompt itself will use the first 2000 characters.
        # This is important for large essays.
        truncated_essay_text = essay_text[:5000] # Adjust as needed based on model context window

        return f"""
        You are an expert writing assistant specializing in {essay_type}s.
        
        Analyze the following essay and provide corrections with tracked changes:
        1. Use <del>text</del> for text that should be deleted
        2. Use <ins>text</ins> for text that should be added
        3. Focus on grammar, style, clarity, and {essay_type.lower()} specific improvements
        4. Maintain the original meaning and structure
        5. Provide specific suggestions for improvement
        
        Return a valid JSON object:
        {{
          "essay_type": "{essay_type}",
          "essay_score": "numerical score out of 100",
          "corrected_essay": "Essay text with <del> and <ins> tags for tracked changes",
          "suggestions": ["specific suggestion 1", "specific suggestion 2", "specific suggestion 3"]
        }}
        
        Essay to analyze:
        \"\"\"
        {truncated_essay_text}
        \"\"\"
        """
    
    def _fallback_analysis(self, text):
        return {
            "overall_score": "75",
            "suggestions": ["Unable to perform detailed analysis due to missing AI model."]
        }

class DocumentComparator:
    @staticmethod
    def tag_word_diff(line1: str, line2: str) -> Tuple[str, str]:
        diff = difflib.ndiff(line1.split(), line2.split())
        tagged_line1, tagged_line2 = [], []
        
        for word in diff:
            if word.startswith('- '):
                tagged_line1.append(f"<del>{word[2:]}</del>")
            elif word.startswith('+ '):
                tagged_line2.append(f"<ins>{word[2:]}</ins>")
            elif word.startswith('  '): # '  ' means unchanged
                word_clean = word[2:]
                tagged_line1.append(word_clean)
                tagged_line2.append(word_clean)
        return ' '.join(tagged_line1), ' '.join(tagged_line2)

# Initialize components
analyzer = EssayAnalyzer()
doc_processor = DocumentProcessor()
comparator = DocumentComparator()


@api_view(['POST'])
@parser_classes([MultiPartParser]) # For file uploads
def analyze_essay_view(request):
    try:
        # request.FILES for uploaded files in Django
        if 'file' not in request.FILES:
            return Response({'error': 'No file uploaded'}, status=status.HTTP_400_BAD_REQUEST)
        
        file = request.FILES['file']
        
        # Django's UploadedFile objects are file-like, but you need to check filename directly
        if not file.name.lower().endswith('.docx'):
            return Response({'error': 'Only .docx files are supported'}, status=status.HTTP_400_BAD_REQUEST)
        
        # Pass the UploadedFile object directly to docx library
        essay_text = doc_processor.extract_text_from_docx(file)
        if not essay_text or not essay_text.strip():
            return Response({'error': 'No text found in the document'}, status=status.HTTP_400_BAD_REQUEST)
        if len(essay_text.strip()) < 50:
            return Response({'error': 'Essay is too short for meaningful analysis'}, status=status.HTTP_400_BAD_REQUEST)

        global model # Access the global model
        if not model:
            return Response({
                'essay_score': '75',
                'essay_type': 'Expository Essay',
                'corrected_essay': essay_text,
                'suggestions': ['AI model unavailable - please check your API key or model configuration'],
            }, status=status.HTTP_503_SERVICE_UNAVAILABLE)

        detected_type = analyzer.classify_essay_type(essay_text)
        
        correction_prompt = analyzer.create_correction_prompt(essay_text, detected_type)
        try:
            correction_response = model.generate_content(correction_prompt)
            corrected_json = analyzer.extract_json(correction_response.text)
        except Exception as e:
            print(f"Error generating content from Gemini: {e}")
            return Response({'error': f'AI processing failed: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        if not corrected_json:
            return Response({'error': 'Essay analysis failed - invalid AI response format'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        return Response({
            'essay_score': corrected_json.get('essay_score', '75'),
            'essay_type': corrected_json.get('essay_type', detected_type),
            'corrected_essay': corrected_json.get('corrected_essay', essay_text),
            'suggestions': corrected_json.get('suggestions', [])
        }, status=status.HTTP_200_OK)

    except Exception as e:
        print(f"Error in analyze_essay_view: {e}")
        return Response({'error': f'Analysis failed: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@parser_classes([JSONParser]) # For JSON data in request body
def change_essay_type_view(request):
    try:
        data = request.data # request.data for DRF parsed JSON
        essay_text = data.get('essay_text')
        target_type = data.get('target_essay_type')
        
        if not essay_text or not target_type:
            return Response({'error': 'essay_text and target_essay_type are required'}, status=status.HTTP_400_BAD_REQUEST)
        if target_type not in ALLOWED_ESSAY_TYPES:
            return Response({'error': f'Invalid essay type. Allowed types: {ALLOWED_ESSAY_TYPES}'}, status=status.HTTP_400_BAD_REQUEST)

        global model
        if not model:
            return Response({
                'essay_score': '75',
                'essay_type': target_type,
                'corrected_essay': essay_text,
                'suggestions': ['AI model unavailable - please check your API key or model configuration'],
            }, status=status.HTTP_503_SERVICE_UNAVAILABLE)

        clean_text = re.sub(r'<[^>]+>', '', essay_text) # Clean existing HTML tags
        
        correction_prompt = analyzer.create_correction_prompt(clean_text, target_type)
        try:
            correction_response = model.generate_content(correction_prompt)
            corrected_json = analyzer.extract_json(correction_response.text)
        except Exception as e:
            print(f"Error generating content from Gemini: {e}")
            return Response({'error': f'AI processing failed: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        if not corrected_json:
            return Response({'error': 'Essay type conversion failed - invalid AI response format'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        return Response({
            'essay_score': corrected_json.get('essay_score', '75'),
            'essay_type': target_type, # Use the target_type, as AI might deviate
            'corrected_essay': corrected_json.get('corrected_essay', essay_text),
            'suggestions': corrected_json.get('suggestions', [])
        }, status=status.HTTP_200_OK)

    except Exception as e:
        print(f"Error in change_essay_type_view: {e}")
        return Response({'error': f'Type conversion failed: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@api_view(['POST'])
@parser_classes([JSONParser])
def download_revision_view(request):
    try:
        data = request.data
        final_text = data.get('final_text', '')
        title = data.get('title', 'Revised Essay')
        
        if not final_text:
            return Response({'error': 'No text provided for download'}, status=status.HTTP_400_BAD_REQUEST)
        
        clean_text = re.sub(r'<[^>]+>', '', final_text) # Clean HTML tags from the text
        
        doc = Document()
        doc.add_heading(title, 0)
        
        paragraphs = clean_text.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # Django's way to send a file from an in-memory BytesIO object
        response = HttpResponse(
            doc_io.getvalue(), # Get the bytes content
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{title.replace(" ", "_")}_final.docx"'
        return response
            
    except Exception as e:
        print(f"Error in download_revision_view: {e}")
        return Response({'error': 'Failed to create download file'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['POST'])
@parser_classes([MultiPartParser])
def compare_essays_view(request):
    try:
        if 'essay1' not in request.FILES or 'essay2' not in request.FILES:
            return Response({'error': 'Please upload both essay files'}, status=status.HTTP_400_BAD_REQUEST)

        file1 = request.FILES['essay1']
        file2 = request.FILES['essay2']

        # Check file extensions
        allowed_extensions = ('.txt', '.docx', '.pdf')
        if not (file1.name.lower().endswith(allowed_extensions) and
                file2.name.lower().endswith(allowed_extensions)):
            return Response({'error': 'Both files must be TXT, DOCX, or PDF format'}, status=status.HTTP_400_BAD_REQUEST)

        # Pass file objects and their names
        lines1 = doc_processor.extract_lines(file1, file1.name)
        file1.seek(0) # Reset file pointer after reading for potential re-reads
        lines2 = doc_processor.extract_lines(file2, file2.name)
        file2.seek(0) # Reset file pointer

        global model
        if not model:
            return Response({'error': 'AI analysis unavailable - please check your API key or model configuration'}, status=status.HTTP_503_SERVICE_UNAVAILABLE)

        full_essay1 = "\n".join(lines1)
        full_essay2 = "\n".join(lines2)
        
        # Truncate essays for prompt to avoid exceeding token limits
        # Remember to adjust these limits based on the model's actual context window and your needs
        truncated_essay1 = full_essay1[:4000] # Example limit
        truncated_essay2 = full_essay2[:4000] # Example limit

        analysis_prompt1 = f"""
        Analyze this essay and provide a brief summary of its characteristics (tone, style, structure, main ideas) in 2-3 sentences:
        
        {truncated_essay1}...
        """
        
        analysis_prompt2 = f"""
        Analyze this essay and provide a brief summary of its characteristics (tone, style, structure, main ideas) in 2-3 sentences:
        
        {truncated_essay2}...
        """
        
        try:
            response1 = model.generate_content(analysis_prompt1)
            essay1_analysis = response1.text.strip()
        except Exception as e:
            essay1_analysis = f"Error analyzing first essay: {str(e)}"
            print(f"Error in essay1 analysis: {e}")
            
        try:
            response2 = model.generate_content(analysis_prompt2)
            essay2_analysis = response2.text.strip()
        except Exception as e:
            essay2_analysis = f"Error analyzing second essay: {str(e)}"
            print(f"Error in essay2 analysis: {e}")

        # Generate summary of key differences
        # Truncate for comparison prompt as well
        truncated_compare_essay1 = full_essay1[:2000] # Example limit
        truncated_compare_essay2 = full_essay2[:2000] # Example limit

        summary_prompt = f"""
        Compare these two essays and identify key differences in tone, structure, ideas, and writing style.
        Provide 3-5 bullet points highlighting the main differences:

        Essay 1: {truncated_compare_essay1}...
        Essay 2: {truncated_compare_essay2}...
        """

        try:
            summary_response = model.generate_content(summary_prompt)
            key_differences = summary_response.text.strip()
        except Exception as e:
            key_differences = f"Error generating comparison: {str(e)}"
            print(f"Error in comparison generation: {e}")

        return Response({
            'status': 'success',
            'draft1_analysis': essay1_analysis,
            'draft2_analysis': essay2_analysis,
            'key_differences': key_differences,
            'file1_name': file1.name, # Django provides name directly
            'file2_name': file2.name,
            'timestamp': datetime.now().isoformat()
        }, status=status.HTTP_200_OK)

    except Exception as e:
        print(f"Error in compare_essays_view: {e}")
        return Response({'error': f'Essay comparison failed: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)